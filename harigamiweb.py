import streamlit as st
import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import warnings
import io
import zipfile
import re

warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

DEFAULT_TEMPLATE_PATH = "harigami.docx"
OUTPUT_DIR = "output_docs"

PLACEHOLDERS = {
    "［10月　19日（水）］": "DATE",
    "［10:00］": "START_TIME",
    "［11:00］": "END_TIME",
    "［物件名］": "NAME",
}

def replace_placeholders_preserve_format(paragraph, replacements):
    full_text = paragraph.text
    should_center = False

    for ph, key in PLACEHOLDERS.items():
        if ph in full_text:
            if key in ["DATE", "START_TIME", "END_TIME"]:
                should_center = True

            for run in paragraph.runs:
                if ph in run.text:
                    original_font_size = run.font.size
                    original_bold = run.font.bold
                    original_italic = run.font.italic
                    original_underline = run.font.underline
                    original_color = run.font.color
                    
                    run.text = run.text.replace(ph, replacements[key])
                    
                    if original_font_size:
                        run.font.size = original_font_size
                    if original_bold is not None:
                        run.font.bold = original_bold
                    if original_italic is not None:
                        run.font.italic = original_italic
                    if original_underline is not None:
                        run.font.underline = original_underline
                    if original_color:
                        run.font.color.rgb = original_color.rgb
                    break
            
            current_text = paragraph.text
            if ph in current_text:
                replace_text_across_runs(paragraph, ph, replacements[key])
    
    if should_center:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def replace_text_across_runs(paragraph, search_text, replace_text):
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if search_text in full_text:
        new_text = full_text.replace(search_text, replace_text)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            for run in paragraph.runs[1:]:
                run.text = ""
            first_run.text = new_text

def replace_placeholders_in_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_preserve_format(paragraph, replacements)

def replace_placeholders_comprehensive(doc, replacements):
    for para in doc.paragraphs:
        if para.text.strip():
            replace_placeholders_preserve_format(para, replacements)
    
    replace_placeholders_in_tables(doc, replacements)
    
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                replace_placeholders_preserve_format(para, replacements)
        
        if section.footer:
            for para in section.footer.paragraphs:
                replace_placeholders_preserve_format(para, replacements)

def process_excel_and_generate_docs(excel_file_buffer, template_source, is_uploaded):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generated_file_paths = []

    try:
        df = pd.read_excel(excel_file_buffer, sheet_name="作業指示書 の一覧", engine='openpyxl')
        total_rows = len(df)
        st.info(f"📊 {total_rows}件のデータを読み込みました。文書生成を開始します...")
        
        progress_text = st.empty()
        progress_bar = st.progress(0)
        processed_count = 0

        for index, row in df.iterrows():
            progress_percent = (index + 1) / total_rows
            progress_bar.progress(progress_percent)
            progress_text.text(f"処理中: {index + 1} / {total_rows} 件完了")

            try:
                if pd.isna(row["物件名"]) or pd.isna(row["予定開始"]) or pd.isna(row["予定終了"]):
                    continue
                
                name = str(row["物件名"]).strip()
                start_dt = pd.to_datetime(row["予定開始"], errors='coerce')
                end_dt = pd.to_datetime(row["予定終了"], errors='coerce')
                if pd.isna(start_dt) or pd.isna(end_dt):
                    continue

                weekdays = {'Mon': '月', 'Tue': '火', 'Wed': '水', 'Thu': '木', 
                            'Fri': '金', 'Sat': '土', 'Sun': '日'}
                weekday = weekdays[start_dt.strftime('%a')]
                date_str = f"{start_dt.month}月{start_dt.day}日（{weekday}）"
                start_str = start_dt.strftime("%H:%M")
                end_str = end_dt.strftime("%H:%M")

            except Exception:
                continue

            replacements = {
                "DATE": date_str,
                "START_TIME": start_str,
                "END_TIME": end_str,
                "NAME": name
            }

            try:
                if is_uploaded:
                    template_source.seek(0)
                    doc = Document(template_source)
                else:
                    doc = Document(template_source)

                replace_placeholders_comprehensive(doc, replacements)

                safe_name = re.sub(r'[^\w\.\-]', '_', name)
                safe_name = re.sub(r'_{2,}', '_', safe_name)
                safe_name = safe_name.strip('_') or "untitled_document"

                output_file_name = f"{safe_name}.docx"
                output_path = os.path.join(OUTPUT_DIR, output_file_name)
                doc.save(output_path)

                if os.path.exists(output_path):
                    generated_file_paths.append(output_path)
                    processed_count += 1

            except Exception:
                continue

        progress_bar.progress(1.0)
        progress_text.text(f"処理完了: {processed_count} / {total_rows} 件完了")
        st.success(f"\n🎉 {processed_count}件の通知文書の生成が完了しました！")
        return generated_file_paths

    except Exception as e:
        st.error(f"❌ エラーが発生しました: {str(e)}")
        return []

# --- Streamlit UI ---
st.set_page_config(
    page_title="貼紙自動生成アプリ",
    page_icon="📄",
    layout="centered"
)

st.title("📄貼紙自動生成アプリ")
st.markdown("""
このアプリは、アップロードされたExcelファイル（作業指示書一覧）の情報に基づいて、
貼紙のWord文書を生成します。
""")

uploaded_file = st.file_uploader(
    "1. Excelファイルを選択してください (.xlsx, .xls)",
    type=["xlsx", "xls"],
    help="「作業指示書 の一覧」シートが含まれるExcelファイルをアップロードしてください。"
)

template_choice = st.radio(
    "2. 使用するテンプレートを選択してください",
    ("デフォルトテンプレートを使用", "テンプレートをアップロードする")
)

if template_choice == "テンプレートをアップロードする":
    uploaded_template = st.file_uploader(
        "2-1. Wordテンプレートファイルをアップロードしてください (.docx)",
        type=["docx"]
    )
    template_ready = uploaded_template is not None
    template_info = uploaded_template
    is_uploaded_template = True
else:
    if not os.path.exists(DEFAULT_TEMPLATE_PATH):
        st.error(f"❌ デフォルトテンプレートが見つかりません: `{DEFAULT_TEMPLATE_PATH}`")
        template_ready = False
        template_info = None
    else:
        st.info(f"✅ デフォルトテンプレート `{DEFAULT_TEMPLATE_PATH}` を使用します。")
        template_ready = True
        template_info = DEFAULT_TEMPLATE_PATH
        is_uploaded_template = False

if uploaded_file and template_ready:
    st.success("Excelファイルとテンプレートが準備できました！")

    if st.button("3. Word文書を生成する"):
        with st.spinner("Word文書を生成中...しばらくお待ちください。"):
            excel_buffer = io.BytesIO(uploaded_file.read())
            if is_uploaded_template:
                template_buffer = io.BytesIO(template_info.read())
                generated_doc_paths = process_excel_and_generate_docs(excel_buffer, template_buffer, True)
            else:
                generated_doc_paths = process_excel_and_generate_docs(excel_buffer, template_info, False)

        if generated_doc_paths:
            st.subheader("🎉 生成された文書をまとめてダウンロード")
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for doc_path in generated_doc_paths:
                    zf.write(doc_path, os.path.basename(doc_path))
            zip_buffer.seek(0)

            st.download_button(
                label="全てのWord文書をZIPでダウンロード",
                data=zip_buffer.getvalue(),
                file_name="generated_word_documents.zip",
                mime="application/zip"
            )

            for doc_path in generated_doc_paths:
                if os.path.exists(doc_path):
                    os.remove(doc_path)
        else:
            st.warning("文書の生成に失敗したか、対象データがありません。")
else:
    st.info("⬆️ Excelファイルとテンプレートファイルをアップロードまたは選択してください。")

st.markdown("---")
st.caption("Powered by Streamlit")
